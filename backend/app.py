import os
import traceback
import PyPDF2
from flask import Flask, request, jsonify, send_file, send_from_directory
from flask_cors import CORS
from docx import Document
import openpyxl
import requests
from dotenv import load_dotenv
import time  # Added for retries

# PDF library fallback
try:
    from fpdf import FPDF
except ImportError:
    from fpdf2 import FPDF as FPDF

app = Flask(__name__, static_folder="../Frontend/dist", static_url_path="")
# Allow only your Render frontend
CORS(app, resources={r"/process": {"origins": "https://query-master-1.onrender.com"}})

# Configuration
load_dotenv()
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB file limit
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'docx', 'txt'}

# Ensure upload directory exists with proper permissions
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'], mode=0o777)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

@app.route("/")
def serve_react():
    return send_from_directory(app.static_folder, "index.html")

def extract_text_from_file(file_path, file_format):
    try:
        if file_format == "pdf":
            with open(file_path, "rb") as file:
                reader = PyPDF2.PdfReader(file)
                text = "\n".join([page.extract_text() or "" for page in reader.pages]).strip()
        elif file_format == "docx":
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs if para.text]).strip()
        elif file_format == "txt":
            with open(file_path, "r", encoding='utf-8') as file:
                text = file.read().strip()
        else:
            raise ValueError(f"Unsupported file format: {file_format}")

        return text or "No extractable text found"
    except Exception as e:
        raise RuntimeError(f"Failed to extract text: {str(e)}")

def generate_answers(questions, context):
    url = "https://chatgpt-42.p.rapidapi.com/chat"
    headers = {
        "x-rapidapi-key": os.getenv("RAPIDAPI_KEY"),
        "x-rapidapi-host": "chatgpt-42.p.rapidapi.com",
        "Content-Type": "application/json"
    }

    answers = []
    batch_size = 10  # Reduced for reliability
    unwanted_phrases = ["if you have more questions", "feel free to ask", "let me know if you need"]
    max_retries = 3

    for i in range(0, len(questions), batch_size):
        batch = questions[i:i + batch_size]
        payload = {
            "messages": [{
                "role": "user",
                "content": f"Context: {context}\n\nQuestions:\n" + "\n".join(batch)
            }],
            "model": "gpt-4o-mini",
            "max_tokens": 1000
        }

        for attempt in range(max_retries):
            try:
                response = requests.post(url, json=payload, headers=headers, timeout=30)
                response.raise_for_status()
                answer = response.json().get("choices", [{}])[0].get("message", {}).get("content", "")

                for phrase in unwanted_phrases:
                    answer = answer.lower().replace(phrase, "").strip()

                answers.extend(filter(None, answer.split("\n")))
                break  # Successful request, break retry loop
            except Exception as e:
                if attempt < max_retries - 1:
                    time.sleep(2)  # Small delay before retry
                else:
                    raise RuntimeError(f"API request failed after retries: {str(e)}")

    return answers or ["No answers generated"]

def save_answers(answers, file_format):
    try:
        if not answers:
            raise ValueError("No answers to save")

        file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"answers.{file_format}")

        if file_format == "txt":
            with open(file_path, "w", encoding='utf-8') as f:
                f.write("\n".join(answers))

        elif file_format == "docx":
            doc = Document()
            for answer in answers:
                doc.add_paragraph(answer)
            doc.save(file_path)

        elif file_format == "xlsx":
            wb = openpyxl.Workbook()
            ws = wb.active
            for i, answer in enumerate(answers):
                ws.cell(row=i+1, column=1, value=answer)
            wb.save(file_path)

        elif file_format == "pdf":
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.cell(200, 10, txt="Generated Answers", ln=True, align="C")
            pdf.ln(10)
            for answer in answers:
                pdf.multi_cell(0, 10, txt=answer)
                pdf.ln(5)
            pdf.output(file_path)

        return file_path
    except Exception as e:
        raise RuntimeError(f"Failed to save {file_format}: {str(e)}")

@app.route("/process", methods=["POST"])
def process_file():
    try:
        if 'file' not in request.files:
            return jsonify({"error": "No file part"}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({"error": "No selected file"}), 400

        if not allowed_file(file.filename):
            return jsonify({"error": "Invalid file type", "allowed": list(app.config['ALLOWED_EXTENSIONS'])}), 400

        # Ensure the upload folder exists before saving
        if not os.path.exists(app.config['UPLOAD_FOLDER']):
            os.makedirs(app.config['UPLOAD_FOLDER'], mode=0o777)

        input_format = request.form.get("input_format", "pdf")
        output_format = request.form.get("output_format", "txt")

        file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(file_path)

        text = extract_text_from_file(file_path, input_format)
        questions = [q.strip() + "?" for q in text.replace("\n", " ").split("?") if q.strip()]

        if not questions:
            os.remove(file_path)  # Cleanup
            return jsonify({"error": "No questions detected"}), 400

        answers = generate_answers(questions, text)
        result_file = save_answers(answers, output_format)

        os.remove(file_path)  # Cleanup after processing

        return jsonify({"success": True, "download_link": f"/download/{os.path.basename(result_file)}"})

    except Exception as e:
        app.logger.error(f"Processing error: {str(e)}\n{traceback.format_exc()}")
        return jsonify({"error": "File processing failed", "details": str(e)}), 500

@app.route("/download/<filename>", methods=["GET"])
def download_file(filename):
    try:
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if not os.path.exists(file_path):
            return jsonify({"error": "File not found"}), 404
        return send_file(file_path, as_attachment=True)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/<path:path>")
def serve_static(path):
    return send_from_directory(app.static_folder, path)

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)



# import os
# import traceback
# import PyPDF2
# from flask import Flask, request, jsonify, send_file, send_from_directory
# from flask_cors import CORS
# from docx import Document
# import openpyxl
# import requests
# from dotenv import load_dotenv
# # from docx.shared import Pt  
# # from docx.oxml import OxmlElement  
# # from docx.oxml.ns import qn

# # PDF library fallback
# try:
#     from fpdf import FPDF
# except ImportError:
#     from fpdf2 import FPDF as FPDF

# app = Flask(__name__, static_folder="../Frontend/dist", static_url_path="")
# # Allow only your Render frontend
# CORS(app, resources={
#     r"/process": {"origins": "https://query-master-1.onrender.com"}
# })

# # Configuration
# load_dotenv()
# app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB file limit
# app.config['UPLOAD_FOLDER'] = 'uploads'
# app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'docx', 'txt'}

# # Ensure upload directory exists with proper permissions
# if not os.path.exists(app.config['UPLOAD_FOLDER']):
#     os.makedirs(app.config['UPLOAD_FOLDER'], mode=0o777)

# def allowed_file(filename):
#     return '.' in filename and \
#            filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

# @app.route("/")
# def serve_react():
#     return send_from_directory(app.static_folder, "index.html")

# def extract_text_from_file(file_path, file_format):
#     try:
#         if file_format == "pdf":
#             with open(file_path, "rb") as file:
#                 reader = PyPDF2.PdfReader(file)
#                 text = ""
#                 for page in reader.pages:
#                     text += page.extract_text() or ""  # Handle None returns
#         elif file_format == "docx":
#             doc = Document(file_path)
#             text = "\n".join([para.text for para in doc.paragraphs if para.text])
#         elif file_format == "txt":
#             with open(file_path, "r", encoding='utf-8') as file:
#                 text = file.read()
#         else:
#             raise ValueError(f"Unsupported file format: {file_format}")
#         return text.strip() or "No extractable text found"
#     except Exception as e:
#         raise RuntimeError(f"Failed to extract text: {str(e)}")

# def generate_answers(questions, context):
#     url = "https://chatgpt-42.p.rapidapi.com/chat"
#     headers = {
#         "x-rapidapi-key": os.getenv("RAPIDAPI_KEY"),
#         "x-rapidapi-host": "chatgpt-42.p.rapidapi.com",
#         "Content-Type": "application/json"
#     }

#     answers = []
#     batch_size = 10  # Reduced for reliability
#     unwanted_phrases = [
#         "if you have more questions",
#         "feel free to ask",
#         "let me know if you need"
#     ]

#     for i in range(0, len(questions), batch_size):
#         try:
#             batch = questions[i:i + batch_size]
#             payload = {
#                 "messages": [{
#                     "role": "user",
#                     "content": f"Context: {context}\n\nQuestions:\n" + "\n".join(batch)
#                 }],
#                 "model": "gpt-4o-mini",
#                 "max_tokens": 1000
#             }
            
#             response = requests.post(url, json=payload, headers=headers, timeout=30)
#             response.raise_for_status()
            
#             answer = response.json().get("choices", [{}])[0].get("message", {}).get("content", "")
#             for phrase in unwanted_phrases:
#                 answer = answer.lower().replace(phrase, "").strip()
            
#             answers.extend(filter(None, answer.split("\n")))
            
#         except Exception as e:
#             raise RuntimeError(f"API request failed: {str(e)}")

#     return answers or ["No answers generated"]



# def save_answers(answers, format):
#     try:
#         if not answers:
#             raise ValueError("No answers to save")
            
#         if format == "txt":
#             file_path = os.path.join(app.config['UPLOAD_FOLDER'], "answers.txt")
#             with open(file_path, "w", encoding='utf-8') as f:
#                 f.write("\n".join(answers))
                
#         elif format == "docx":
#             file_path = os.path.join(app.config['UPLOAD_FOLDER'], "answers.docx")
#             doc = Document()
#             for answer in answers:
#                 doc.add_paragraph(answer)
#             doc.save(file_path)
  

            
#         elif format == "xls":
#             file_path = os.path.join(app.config['UPLOAD_FOLDER'], "answers.xlsx")
#             wb = openpyxl.Workbook()
#             ws = wb.active
#             for i, answer in enumerate(answers):
#                 ws.cell(row=i+1, column=1, value=answer)
#             wb.save(file_path)
            
#         elif format == "pdf":
#             file_path = os.path.join(app.config['UPLOAD_FOLDER'], "answers.pdf")
#             pdf = FPDF()
#             pdf.add_page()
#             pdf.set_font("Arial", size=12)
#             pdf.cell(200, 10, txt="Generated Answers", ln=True, align="C")
#             pdf.ln(10)
            
#             for answer in answers:
#                 pdf.multi_cell(0, 10, txt=answer)
#                 pdf.ln(5)
                
#             pdf.output(file_path)
            
#         return file_path
        
#     except Exception as e:
#         raise RuntimeError(f"Failed to save {format}: {str(e)}")



# @app.route("/process", methods=["POST"])
# def process_file():
#     try:
#         # Validate request
#         if 'file' not in request.files:
#             return jsonify({"error": "No file part"}), 400
            
#         file = request.files['file']
#         if file.filename == '':
#             return jsonify({"error": "No selected file"}), 400
            
#         if not allowed_file(file.filename):
#             return jsonify({
#                 "error": "Invalid file type",
#                 "allowed": list(app.config['ALLOWED_EXTENSIONS'])
#             }), 400

#         # Process file
#         input_format = request.form.get("input_format", "pdf")
#         output_format = request.form.get("output_format", "txt")
        
#         file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
#         file.save(file_path)
        
#         text = extract_text_from_file(file_path, input_format)
#         questions = [f"{q.strip()}?" for q in text.split("?") if q.strip()]
        
#         if not questions:
#             return jsonify({"error": "No questions detected"}), 400
            
#         answers = generate_answers(questions, text)
#         result_file = save_answers(answers, output_format)
        
#         return jsonify({
#             "success": True,
#             "download_link": f"/download/{os.path.basename(result_file)}"
#         })
        
#     except Exception as e:
#         app.logger.error(f"Processing error: {str(e)}\n{traceback.format_exc()}")
#         return jsonify({
#             "error": "File processing failed",
#             "details": str(e)
#         }), 500

# @app.route("/download/<filename>", methods=["GET"])
# def download_file(filename):
#     try:
#         file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#         if not os.path.exists(file_path):
#             return jsonify({"error": "File not found"}), 404
#         return send_file(file_path, as_attachment=True)
#     except Exception as e:
#         return jsonify({"error": str(e)}), 500

# @app.route("/<path:path>")
# def serve_static(path):
#     return send_from_directory(app.static_folder, path)

# # Updated app.run() for Render
# if __name__ == "__main__":
#     port = int(os.environ.get("PORT", 5000))
#     app.run(host="0.0.0.0", port=port, debug=False)















# from flask import Flask, request, jsonify, send_file
# from flask_cors import CORS
# import os
# import PyPDF2
# from docx import Document
# import openpyxl
# import requests
# from dotenv import load_dotenv
# from fpdf import FPDF

# app = Flask(__name__)
# CORS(app)

# # Load environment variables from .env file
# load_dotenv()

# # Temporary storage for files
# UPLOAD_FOLDER = "uploads"
# if not os.path.exists(UPLOAD_FOLDER):
#     os.makedirs(UPLOAD_FOLDER)

# # Root route
# @app.route("/")
# def home():
#     return "QueryMaster Flask Backend is running!"

# # Extract text from file
# def extract_text_from_file(file_path, file_format):
#     try:
#         if file_format == "pdf":
#             with open(file_path, "rb") as file:
#                 reader = PyPDF2.PdfReader(file)
#                 text = ""
#                 for page in reader.pages:
#                     text += page.extract_text()
#         elif file_format == "docx":
#             doc = Document(file_path)
#             text = "\n".join([para.text for para in doc.paragraphs])
#         elif file_format == "txt":
#             with open(file_path, "r") as file:
#                 text = file.read()
#         else:
#             raise Exception("Unsupported file format")
#         return text
#     except Exception as e:
#         raise Exception(f"Failed to extract text from file: {str(e)}")

# # Generate answers using OpenAI API (via RapidAPI)
# def generate_answers(questions, context):
#     url = "https://chatgpt-42.p.rapidapi.com/chat"
#     headers = {
#         "x-rapidapi-key": os.getenv("RAPIDAPI_KEY"),  # Load API key from .env
#         "x-rapidapi-host": "chatgpt-42.p.rapidapi.com",
#         "Content-Type": "application/json"
#     }

#     answers = []
#     batch_size = 15  # Number of questions to send in each API call
#     for i in range(0, len(questions), batch_size):
#         batch = questions[i:i + batch_size]
#         try:
#             # Combine all questions in the batch into a single prompt
#             combined_prompt = "Context: " + context + "\n\nQuestions:\n" + "\n".join(batch)
#             payload = {
#                 "messages": [
#                     {
#                         "role": "user",
#                         "content": combined_prompt
#                     }
#                 ],
#                 "model": "gpt-4o-mini"
#             }
#             response = requests.post(url, json=payload, headers=headers)
#             response.raise_for_status()
#             # Assuming the API returns a single response for all questions
#             answer = response.json()["choices"][0]["message"]["content"]
#             # Split the answer into individual answers
#             individual_answers = answer.split("\n")
#             answers.extend(individual_answers)
#         except Exception as e:
#             raise Exception(f"API error: {str(e)}")
#     return answers

# # Save answers to a file
# def save_answers(answers, format):
#     try:
#         if format == "txt":
#             file_path = os.path.join(UPLOAD_FOLDER, "answers.txt")
#             with open(file_path, "w") as file:
#                 for answer in answers:
#                     file.write(answer + "\n")
#         elif format == "docx":
#             file_path = os.path.join(UPLOAD_FOLDER, "answers.docx")
#             doc = Document()
#             for answer in answers:
#                 doc.add_paragraph(answer)
#             doc.save(file_path)
#         elif format == "xls":
#             file_path = os.path.join(UPLOAD_FOLDER, "answers.xlsx")
#             wb = openpyxl.Workbook()
#             ws = wb.active
#             for i, answer in enumerate(answers):
#                 ws.cell(row=i+1, column=1, value=answer)
#             wb.save(file_path)
#         elif format == "pdf":
#             file_path = os.path.join(UPLOAD_FOLDER, "answers.pdf")
#             # Create a PDF using fpdf
#             pdf = FPDF()
#             pdf.add_page()
#             pdf.set_font("Arial", size=12)  # Set font and size

#             # Add a title
#             pdf.cell(200, 10, txt="Generated Answers", ln=True, align="C")
#             pdf.ln(10)  # Add some space after the title

#             # Write each answer to the PDF
#             for answer in answers:
#                 pdf.multi_cell(0, 10, txt=answer)  # Automatically handles line breaks
#                 pdf.ln(5)  # Add some space between answers

#             # Save the PDF
#             pdf.output(file_path)
#         return file_path
#     except Exception as e:
#         raise Exception(f"Failed to save answers: {str(e)}")

# # API endpoint to process file and generate answers
# @app.route("/process", methods=["POST"])
# def process_file():
#     try:
#         if "file" not in request.files:
#             return jsonify({"error": "No file uploaded"}), 400

#         file = request.files["file"]
#         input_format = request.form.get("input_format", "pdf")
#         output_format = request.form.get("output_format", "txt")

#         # Save the uploaded file
#         file_path = os.path.join(UPLOAD_FOLDER, file.filename)
#         file.save(file_path)

#         # Extract text from the file
#         text = extract_text_from_file(file_path, input_format)

#         # Extract questions from the text
#         questions = text.split("?")  # Simple question extraction (can be improved)
#         questions = [q.strip() + "?" for q in questions if q.strip()]

#         if not questions:
#             return jsonify({"error": "No questions found in the file"}), 400

#         # Generate answers using OpenAI API
#         answers = generate_answers(questions, text)

#         # Save answers to the selected output format
#         file_path = save_answers(answers, output_format)

#         # Return the download link
#         return jsonify({"download_link": f"/download/{os.path.basename(file_path)}"})
#     except Exception as e:
#         return jsonify({"error": str(e)}), 500

# # API endpoint to download the file
# @app.route("/download/<filename>", methods=["GET"])
# def download_file(filename):
#     file_path = os.path.join(UPLOAD_FOLDER, filename)
#     return send_file(file_path, as_attachment=True)

# if __name__ == "__main__":
#     app.run(debug=True)















# from flask import Flask, request, jsonify, send_file
# from flask_cors import CORS
# import os
# import PyPDF2
# from docx import Document
# import openpyxl
# import requests
# from fpdf import FPDF
# from dotenv import load_dotenv

# app = Flask(__name__)
# CORS(app)

# # Load environment variables from .env file
# load_dotenv()

# # Temporary storage for files
# UPLOAD_FOLDER = "uploads"
# if not os.path.exists(UPLOAD_FOLDER):
#     os.makedirs(UPLOAD_FOLDER)

# # Root route
# @app.route("/")
# def home():
#     return "QueryMaster Flask Backend is running!"

# # Extract text from file
# def extract_text_from_file(file_path, file_format):
#     try:
#         if file_format == "pdf":
#             with open(file_path, "rb") as file:
#                 reader = PyPDF2.PdfReader(file)
#                 text = ""
#                 for page in reader.pages:
#                     text += page.extract_text()
#         elif file_format == "docx":
#             doc = Document(file_path)
#             text = "\n".join([para.text for para in doc.paragraphs])
#         elif file_format == "txt":
#             with open(file_path, "r") as file:
#                 text = file.read()
#         else:
#             raise Exception("Unsupported file format")
#         return text
#     except Exception as e:
#         raise Exception(f"Failed to extract text from file: {str(e)}")

# # Generate answers using the new API
# def generate_answers(questions, context):
#     url = "https://chatgpt-42.p.rapidapi.com/chat"
#     headers = {
#         "x-rapidapi-key": os.getenv("RAPIDAPI_KEY"),  # Load API key from .env
#         "x-rapidapi-host": "chatgpt-42.p.rapidapi.com",
#         "Content-Type": "application/json"
#     }

#     answers = []
#     for question in questions:
#         try:
#             payload = {
#                 "messages": [
#                     {
#                         "role": "user",
#                         "content": f"Context: {context}\n\nQuestion: {question}"
#                     }
#                 ],
#                 "model": "gpt-4o-mini"
#             }
#             response = requests.post(url, json=payload, headers=headers)
#             response.raise_for_status()
#             answer = response.json()["choices"][0]["message"]["content"]
#             answers.append(answer)
#         except Exception as e:
#             raise Exception(f"API error: {str(e)}")
#     return answers

# # Save answers to a file
# def save_answers(answers, format):
#     try:
#         if format == "txt":
#             file_path = os.path.join(UPLOAD_FOLDER, "answers.txt")
#             with open(file_path, "w") as file:
#                 for answer in answers:
#                     file.write(answer + "\n")
#         elif format == "docx":
#             file_path = os.path.join(UPLOAD_FOLDER, "answers.docx")
#             doc = Document()
#             for answer in answers:
#                 doc.add_paragraph(answer)
#             doc.save(file_path)
#         elif format == "xls":
#             file_path = os.path.join(UPLOAD_FOLDER, "answers.xlsx")
#             wb = openpyxl.Workbook()
#             ws = wb.active
#             for i, answer in enumerate(answers):
#                 ws.cell(row=i+1, column=1, value=answer)
#             wb.save(file_path)
#         elif format == "pdf":
#            file_path = os.path.join(UPLOAD_FOLDER, "answers.pdf")
#             # Create a PDF using fpdf
#             pdf = FPDF()
#             pdf.add_page()
#             pdf.set_font("Arial", size=12)  # Set font and size

#             # Add a title
#             pdf.cell(200, 10, txt="Generated Answers", ln=True, align="C")
#             pdf.ln(10)  # Add some space after the title

#             # Write each answer to the PDF
#             for answer in answers:
#                 pdf.multi_cell(0, 10, txt=answer)  # Automatically handles line breaks
#                 pdf.ln(5)  # Add some space between answers

#             # Save the PDF
#             pdf.output(file_path)
#         return file_path
#     except Exception as e:
#         raise Exception(f"Failed to save answers: {str(e)}")

# # API endpoint to process file and generate answers
# @app.route("/process", methods=["POST"])
# def process_file():
#     try:
#         if "file" not in request.files:
#             return jsonify({"error": "No file uploaded"}), 400

#         file = request.files["file"]
#         input_format = request.form.get("input_format", "pdf")
#         output_format = request.form.get("output_format", "txt")

#         # Save the uploaded file
#         file_path = os.path.join(UPLOAD_FOLDER, file.filename)
#         file.save(file_path)

#         # Extract text from the file
#         text = extract_text_from_file(file_path, input_format)

#         # Extract questions from the text
#         questions = text.split("?")  # Simple question extraction (can be improved)
#         questions = [q.strip() + "?" for q in questions if q.strip()]

#         if not questions:
#             return jsonify({"error": "No questions found in the file"}), 400

#         # Generate answers using the new API
#         answers = generate_answers(questions, text)

#         # Save answers to the selected output format
#         file_path = save_answers(answers, output_format)

#         # Return the download link
#         return jsonify({"download_link": f"/download/{os.path.basename(file_path)}"})
#     except Exception as e:
#         return jsonify({"error": str(e)}), 500

# # API endpoint to download the file
# @app.route("/download/<filename>", methods=["GET"])
# def download_file(filename):
#     file_path = os.path.join(UPLOAD_FOLDER, filename)
#     return send_file(file_path, as_attachment=True)

# if __name__ == "__main__":
#     app.run(debug=True)























# from flask import Flask, request, jsonify, send_file
# from flask_cors import CORS
# import os
# import PyPDF2
# from docx import Document
# import openpyxl
# import requests
# from dotenv import load_dotenv

# app = Flask(__name__)
# CORS(app)

# # Load environment variables from .env file
# load_dotenv()

# # Temporary storage for files
# UPLOAD_FOLDER = "uploads"
# if not os.path.exists(UPLOAD_FOLDER):
#     os.makedirs(UPLOAD_FOLDER)

# # Root route
# @app.route("/")
# def home():
#     return "QueryMaster Flask Backend is running!"

# # Extract text from file
# def extract_text_from_file(file_path, file_format):
#     try:
#         if file_format == "pdf":
#             with open(file_path, "rb") as file:
#                 reader = PyPDF2.PdfReader(file)
#                 text = ""
#                 for page in reader.pages:
#                     text += page.extract_text()
#         elif file_format == "docx":
#             doc = Document(file_path)
#             text = "\n".join([para.text for para in doc.paragraphs])
#         elif file_format == "txt":
#             with open(file_path, "r") as file:
#                 text = file.read()
#         else:
#             raise Exception("Unsupported file format")
#         return text
#     except Exception as e:
#         raise Exception(f"Failed to extract text from file: {str(e)}")

# # Generate answers using OpenAI API (via RapidAPI)
# def generate_answers(questions, context):
#     url = "https://cheapest-gpt-4-turbo-gpt-4-vision-chatgpt-openai-ai-api.p.rapidapi.com/v1/chat/completions"
#     headers = {
#         "x-rapidapi-key": os.getenv("RAPIDAPI_KEY"),  # Load API key from .env
#         "x-rapidapi-host": "cheapest-gpt-4-turbo-gpt-4-vision-chatgpt-openai-ai-api.p.rapidapi.com",
#         "Content-Type": "application/json"
#     }

#     answers = []
#     for question in questions:
#         try:
#             payload = {
#                 "messages": [
#                     {
#                         "role": "user",
#                         "content": f"Context: {context}\n\nQuestion: {question}"
#                     }
#                 ],
#                 "model": "gpt-4o",
#                 "max_tokens": 100,
#                 "temperature": 0.7
#             }
#             response = requests.post(url, json=payload, headers=headers)
#             response.raise_for_status()
#             answer = response.json()["choices"][0]["message"]["content"]
#             answers.append(answer)
#         except Exception as e:
#             raise Exception(f"OpenAI API error: {str(e)}")
#     return answers

# # Save answers to a file
# def save_answers(answers, format):
#     try:
#         if format == "txt":
#             file_path = os.path.join(UPLOAD_FOLDER, "answers.txt")
#             with open(file_path, "w") as file:
#                 for answer in answers:
#                     file.write(answer + "\n")
#         elif format == "docx":
#             file_path = os.path.join(UPLOAD_FOLDER, "answers.docx")
#             doc = Document()
#             for answer in answers:
#                 doc.add_paragraph(answer)
#             doc.save(file_path)
#         elif format == "xls":
#             file_path = os.path.join(UPLOAD_FOLDER, "answers.xlsx")
#             wb = openpyxl.Workbook()
#             ws = wb.active
#             for i, answer in enumerate(answers):
#                 ws.cell(row=i+1, column=1, value=answer)
#             wb.save(file_path)
#         elif format == "pdf":
#             file_path = os.path.join(UPLOAD_FOLDER, "answers.pdf")
#             # Use a library like ReportLab for PDF generation (not implemented here)
#         return file_path
#     except Exception as e:
#         raise Exception(f"Failed to save answers: {str(e)}")

# # API endpoint to process file and generate answers
# @app.route("/process", methods=["POST"])
# def process_file():
#     try:
#         if "file" not in request.files:
#             return jsonify({"error": "No file uploaded"}), 400

#         file = request.files["file"]
#         input_format = request.form.get("input_format", "pdf")
#         output_format = request.form.get("output_format", "txt")

#         # Save the uploaded file
#         file_path = os.path.join(UPLOAD_FOLDER, file.filename)
#         file.save(file_path)

#         # Extract text from the file
#         text = extract_text_from_file(file_path, input_format)

#         # Extract questions from the text
#         questions = text.split("?")  # Simple question extraction (can be improved)
#         questions = [q.strip() + "?" for q in questions if q.strip()]

#         if not questions:
#             return jsonify({"error": "No questions found in the file"}), 400

#         # Generate answers using OpenAI API
#         answers = generate_answers(questions, text)

#         # Save answers to the selected output format
#         file_path = save_answers(answers, output_format)

#         # Return the download link
#         return jsonify({"download_link": f"/download/{os.path.basename(file_path)}"})
#     except Exception as e:
#         return jsonify({"error": str(e)}), 500

# # API endpoint to download the file
# @app.route("/download/<filename>", methods=["GET"])
# def download_file(filename):
#     file_path = os.path.join(UPLOAD_FOLDER, filename)
#     return send_file(file_path, as_attachment=True)

# if __name__ == "__main__":
#     app.run(debug=True)